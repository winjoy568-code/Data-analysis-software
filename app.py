import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import time

# ==========================================
# 0. 系統設定
# ==========================================
st.set_page_config(page_title="生產效能智慧分析系統 Pro", layout="centered")

st.markdown("""
    <style>
    .main { background-color: #ffffff; }
    html, body, [class*="css"] { font-family: Arial, sans-serif; color: #000000; }
    
    h1 { color: #000000; font-weight: 900; font-size: 2.4em; text-align: center; border-bottom: 4px solid #2c3e50; padding-bottom: 15px; margin-bottom: 30px; }
    h2 { color: #1a5276; border-left: 7px solid #1a5276; padding-left: 12px; margin-top: 40px; font-size: 1.6em; font-weight: bold; background-color: #f8f9fa; padding-top: 5px; padding-bottom: 5px; }
    h3 { color: #2e4053; margin-top: 25px; font-size: 1.3em; font-weight: 700; }
    
    p, li, .stMarkdown { font-size: 16px !important; line-height: 1.7 !important; color: #212f3d !important; }
    
    .insight-box { border: 1px solid #d6eaf8; background-color: #ebf5fb; padding: 15px; border-radius: 5px; margin-top: 10px; margin-bottom: 20px; }
    .chart-desc { font-size: 15px; color: #555; background-color: #f9f9f9; padding: 15px; border-left: 4px solid #bdc3c7; margin-bottom: 30px; margin-top: 0px; }
    
    thead tr th:first-child {display:none} tbody th {display:none}
    div.stButton > button { width: 100%; height: 3em; font-weight: bold; font-size: 16px; }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 1. Helper Functions
# ==========================================
def md_to_html(text):
    if not isinstance(text, str): return str(text)
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = text.replace('\n', '<br>')
    return text

def clean_text_for_word(text):
    if not isinstance(text, str): return str(text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'<b>(.*?)</b>', r'\1', text)
    text = re.sub(r'<br>', '\n', text)
    text = re.sub(r'🔴|🟡|🟢', '', text)
    return text.strip()

# ==========================================
# 2. Data Engine
# ==========================================
class DataEngine:
    @staticmethod
    def clean_and_process(df_raw, params):
        try:
            df = df_raw.copy()
            rename_map = {"用電量(kWh)": "耗電量", "產量(雙)": "產量", "OEE(%)": "OEE_RAW", "設備": "機台編號", "機台": "機台編號"}
            for user_col, sys_col in rename_map.items():
                if user_col in df.columns: df = df.rename(columns={user_col: sys_col})
                
            required_cols = ["機台編號", "耗電量", "產量", "OEE_RAW"]
            if not all(col in df.columns for col in required_cols):
                return None, None, f"資料表缺少必要欄位: {[c for c in required_cols if c not in df.columns]}"
            
            if "日期" in df.columns:
                df["日期"] = pd.to_datetime(df["日期"], errors='coerce').dt.date
                if df["日期"].isnull().any(): return None, None, "日期格式錯誤"

            if "廠別" not in df.columns: df["廠別"] = "匯入廠區"

            df["OEE"] = df["OEE_RAW"].apply(lambda x: x / 100.0 if x > 1.0 else x)
            df["單位能耗"] = df.apply(lambda row: row["耗電量"] / row["產量"] if row["產量"] > 0 else 0, axis=1)
            
            valid_energies = df[df["單位能耗"] > 0]["單位能耗"]
            best_energy = valid_energies.min() if not valid_energies.empty else 0
            
            df["能源損失"] = df.apply(lambda row: max(0, (row["單位能耗"] - best_energy) * row["產量"] * params['elec_price']), axis=1)
            df["產能損失機會成本"] = df.apply(
                lambda row: ((params['target_oee']/100 - row["OEE"]) / row["OEE"] * row["產量"] * params['product_margin']) 
                if 0 < row["OEE"] < params['target_oee']/100 else 0, axis=1
            )
            df["總損失"] = df["能源損失"] + df["產能損失機會成本"]
            
            group_col = "廠別" if df["廠別"].nunique() > 1 else "機台編號"
            analysis_scope = "跨廠區分析" if group_col == "廠別" else "單廠設備分析"
            
            summary_agg = df.groupby(group_col).agg({
                "OEE": "mean", "產量": "sum", "耗電量": "sum", 
                "能源損失": "sum", "產能損失機會成本": "sum", "總損失": "sum"
            }).reset_index()
            
            summary_agg["平均單位能耗"] = summary_agg.apply(
                lambda row: row["耗電量"] / row["產量"] if row["產量"] > 0 else 0, axis=1
            )
            summary_agg = summary_agg.sort_values("OEE", ascending=False)
            
            return df, summary_agg, analysis_scope
        except Exception as e:
            return None, None, str(e)

# ==========================================
# 3. Insight Engine
# ==========================================
class InsightEngine:
    @staticmethod
    def generate_narrative(df, summary_agg, group_col, params):
        texts = {}
        target_oee = params['target_oee'] / 100.0
        margin = params['product_margin']
        
        avg_oee = df["OEE"].mean()
        total_loss = df["總損失"].sum()
        best_m = summary_agg.iloc[0]
        worst_m = summary_agg.iloc[-1]
        
        texts['kpi_summary'] = f"本次分析區間內，整體平均 OEE 為 **{avg_oee:.1%}**。其中 **{best_m[group_col]}** 表現最佳，為全廠標竿；而 **{worst_m[group_col]}** 效率敬陪末座，是造成全廠 **NT$ {total_loss:,.0f}** 潛在損失的主要原因。"
        
        eff_gap_pct = 0
        multiplier_msg = ""
        if best_m['平均單位能耗'] > 0 and worst_m['平均單位能耗'] > 0:
            eff_gap_pct = ((worst_m['平均單位能耗'] - best_m['平均單位能耗']) / best_m['平均單位能耗']) * 100
            multiplier = worst_m['平均單位能耗'] / best_m['平均單位能耗']
            multiplier_msg = f"換算下來，**{worst_m[group_col]}** 的耗能是標竿機台的 **{multiplier:.1f} 倍**。"

        texts['benchmark_analysis'] = f"""
        * **標竿設備 ({best_m[group_col]})**：表現最佳，平均 OEE 達 **{best_m['OEE']:.1%}**，單位能耗最低 ({best_m['平均單位能耗']:.5f} kWh/雙)。
        * **瓶頸設備 ({worst_m[group_col]})**：表現最弱，單位生產成本比標竿高出 **{eff_gap_pct:.1f}%**。{multiplier_msg}
        """
        
        texts['rank_desc'] = f"此圖表顯示各設備的綜合實力排名。數據顯示 **{best_m[group_col]}** 位於頂端，顯示其生產效率最優；反之 **{worst_m[group_col]}** 位於底部，建議優先檢討其作業流程。"
        
        # 更新：移除顏色的硬性描述，改為通用描述
        texts['dual_desc'] = "此圖對比了各機台的「產出量 (柱狀)」與「耗電量 (折線)」。正常的生產模式應為「高產出伴隨高耗電」。若發現某設備產出極低，但耗電量曲線卻未等比例下降，即代表存在無效能耗。"
        
        texts['pie_desc'] = "此圖呈現各設備的總用電量佔比。若非主力生產設備卻佔據過高的用電比例，可能代表設備存在漏電、馬達老化或長時間待機未關機的問題。"
        texts['unit_desc'] = f"此圖比較生產每一雙鞋的電力成本。**{best_m[group_col]}** 的柱狀最短，代表能源轉換效率最高；數值過高者建議檢查傳動系統阻力或加熱系統保溫效果。"

        potential_prod = 0
        if worst_m['OEE'] > 0:
            potential_prod = (best_m['OEE'] - worst_m['OEE']) / worst_m['OEE'] * worst_m['產量']
        potential_rev = potential_prod * margin
        texts['opportunity_analysis'] = f"若能將 **{worst_m[group_col]}** 的效率提升至標竿水準，預計本期間可額外生產 **{potential_prod:,.0f} 雙**，相當於挽回 **NT$ {potential_rev:,.0f}** 的營收損失。"

        cv_text = "數據量不足以計算波動率。"
        if len(df) > 1:
            cv_series = df.groupby(group_col)["OEE"].std() / df.groupby(group_col)["OEE"].mean()
            most_stable = cv_series.idxmin()
            most_unstable = cv_series.idxmax()
            cv_text = f"**{most_stable}** 生產節奏最穩定 (CV最低)；**{most_unstable}** 波動最大，顯示製程或人員操作存在變異。"
        texts['stability_analysis'] = cv_text
        texts['cv_desc'] = "變異係數 (CV) 用於衡量生產穩定度。數值越低代表品質與產出越穩定可控；數值過高則代表生產過程極不穩定。"
        texts['scatter_desc'] = "此矩陣圖用於檢視效率與能耗的關聯。**右下角** (高OEE、低能耗) 為理想落點。若數據點落於 **左上角** (低OEE、高能耗)，通常代表設備處於「空轉浪費」狀態。"

        crit_list, avg_list, good_list = [], [], []
        for _, row in summary_agg.iterrows():
            name = row[group_col]
            if row['OEE'] >= target_oee: good_list.append(name)
            elif row['OEE'] >= 0.7: avg_list.append(name)
            else: crit_list.append(name)
            
        action_text = ""
        if crit_list: action_text += f"🔴 **優先改善**：{', '.join(crit_list)}。OEE 低於 70%，請檢查待機未關機狀況。\n\n"
        if avg_list: action_text += f"🟡 **效能提升**：{', '.join(avg_list)}。表現平穩，建議微調參數提升稼動率。\n\n"
        if good_list: action_text += f"🟢 **標竿管理**：{', '.join(good_list)}。運作優異，建議標準化SOP。"
        texts['action_plan'] = action_text
        return texts

# ==========================================
# 3. Viz Engine (視覺化中心) - 多機台優化版
# ==========================================
class VizEngine:
    @staticmethod
    def _common_layout():
        return dict(
            plot_bgcolor='white',
            font=dict(family='Arial, sans-serif', color='black', size=12),
            xaxis=dict(showgrid=True, gridcolor='#f0f0f0'),
            yaxis=dict(showgrid=True, gridcolor='#f0f0f0'),
            margin=dict(l=40, r=40, t=40, b=40)
        )

    @staticmethod
    def create_rank_chart(summary_agg, group_col):
        fig = px.bar(
            summary_agg.sort_values("OEE", ascending=True),
            x="OEE", y=group_col, orientation='h', text="OEE",
            title="綜合實力排名 (依 OEE 排序)"
        )
        fig.update_traces(marker_color='#2E86C1', texttemplate='%{text:.1%}', textposition='outside')
        fig.update_layout(VizEngine._common_layout())
        fig.update_layout(xaxis=dict(range=[0, summary_agg['OEE'].max() * 1.25])) 
        return fig

    @staticmethod
    def create_cv_chart(df, group_col):
        try:
            cv_data = df.groupby(group_col)["OEE"].agg(['mean', 'std'])
            cv_data['CV'] = (cv_data['std'] / cv_data['mean']) * 100
            cv_data = cv_data.fillna(0).reset_index()
            fig = px.bar(cv_data, x=group_col, y="CV", text="CV", title="生產穩定度 (CV變異係數)")
            fig.update_traces(marker_color='#C0392B', texttemplate='%{text:.1f}%', textposition='outside')
            fig.update_layout(VizEngine._common_layout())
            return fig
        except: return go.Figure()

    @staticmethod
    def create_scatter_chart(df, group_col):
        try:
            fig = px.scatter(
                df, x="OEE", y="單位能耗", color=group_col, size="產量",
                title="效率 vs 能耗 關聯分析",
                color_discrete_sequence=px.colors.qualitative.Set1
            )
            fig.update_layout(VizEngine._common_layout())
            return fig
        except: return go.Figure()

    @staticmethod
    def create_dual_axis_chart(df, group_col):
        try:
            # 重構：支援多機台並列顯示
            fig = go.Figure()
            
            # 取得所有機台並分配顏色
            machines = df[group_col].unique()
            colors = px.colors.qualitative.Plotly
            
            for i, machine in enumerate(machines):
                machine_data = df[df[group_col] == machine].sort_values("日期")
                color = colors[i % len(colors)]
                
                # 產量 (Bar)
                fig.add_trace(go.Bar(
                    x=machine_data["日期"], 
                    y=machine_data["產量"], 
                    name=f"{machine} 產量",
                    marker_color=color,
                    opacity=0.6
                ))
                
                # 耗電 (Line)
                fig.add_trace(go.Scatter(
                    x=machine_data["日期"], 
                    y=machine_data["耗電量"], 
                    name=f"{machine} 耗電",
                    yaxis="y2",
                    mode="lines+markers",
                    line=dict(color=color, width=3)
                ))

            layout = VizEngine._common_layout()
            layout.update(dict(
                title="各機台產量與耗電量趨勢對比",
                yaxis=dict(title="產量 (雙)"),
                yaxis2=dict(title="耗電量 (kWh)", overlaying="y", side="right", showgrid=False),
                xaxis=dict(title="日期", tickformat="%Y-%m-%d"),
                barmode='group', # 關鍵：讓 Bar 並排顯示
                legend=dict(orientation="h", y=1.1)
            ))
            fig.update_layout(layout)
            return fig
        except Exception as e:
            return go.Figure()

    @staticmethod
    def create_pie_chart(summary_agg, group_col):
        try:
            fig = px.pie(summary_agg, values="耗電量", names=group_col, hole=0.4, title="總耗電量佔比")
            fig.update_traces(textinfo='percent+label', textfont=dict(size=14, color='black'), marker=dict(colors=px.colors.qualitative.Safe))
            fig.update_layout(VizEngine._common_layout())
            return fig
        except: return go.Figure()

    @staticmethod
    def create_unit_energy_chart(summary_agg, group_col):
        try:
            sorted_agg = summary_agg.sort_values("平均單位能耗")
            fig = px.bar(
                sorted_agg, x=group_col, y="平均單位能耗", text="平均單位能耗",
                title="平均單位能耗 (越低越好)"
            )
            fig.update_traces(marker_color='#145a32', texttemplate='%{text:.5f}', textposition='outside')
            layout = VizEngine._common_layout()
            layout.update(yaxis=dict(range=[0, sorted_agg['平均單位能耗'].max() * 1.2]))
            fig.update_layout(layout)
            return fig
        except: return go.Figure()

# ==========================================
# 4. Report Engine
# ==========================================
class ReportEngine:
    @staticmethod
    def generate_docx(df, summary_agg, texts, figures, analysis_scope):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        head = doc.add_heading('生產效能診斷分析報告', 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"分析範圍：{clean_text_for_word(analysis_scope)}")
        doc.add_paragraph(f"期間：{df['日期'].min()} ~ {df['日期'].max()}")
        doc.add_paragraph("-" * 60)
        
        doc.add_heading('1. 總體績效概覽', level=1)
        doc.add_paragraph(clean_text_for_word(texts['kpi_summary']))
        
        table = doc.add_table(rows=1, cols=len(summary_agg.columns))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, col in enumerate(summary_agg.columns): hdr[i].text = str(col)
        
        for _, row in summary_agg.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                col_name = summary_agg.columns[i]
                if "OEE" in col_name: cells[i].text = f"{val:.1%}"
                elif "能耗" in col_name: cells[i].text = f"{val:.5f}"
                elif "損失" in col_name or "產量" in col_name: cells[i].text = f"{val:,.0f}"
                elif isinstance(val, float): cells[i].text = f"{val:.1f}"
                else: cells[i].text = str(val)
        
        doc.add_heading('2. 深度診斷分析', level=1)
        doc.add_paragraph(clean_text_for_word(texts['benchmark_analysis']))
        doc.add_paragraph(clean_text_for_word(texts['opportunity_analysis']))
        
        def add_chart_with_desc(key, title, desc_key):
            doc.add_heading(title, level=2)
            if key in figures:
                try:
                    img = figures[key].to_image(format="png", width=800, height=400, scale=1.5)
                    doc.add_picture(BytesIO(img), width=Inches(6.0))
                except: doc.add_paragraph("[圖表無法自動生成，請參考網頁版]")
            if desc_key in texts:
                doc.add_paragraph(clean_text_for_word(texts[desc_key]))

        add_chart_with_desc('rank', '綜合實力排名', 'rank_desc')
        add_chart_with_desc('dual', '產量與能耗趨勢', 'dual_desc')
        
        doc.add_heading('3. 電力耗能分析', level=1)
        add_chart_with_desc('pie', '總耗電量佔比', 'pie_desc')
        add_chart_with_desc('unit', '平均單位能耗', 'unit_desc')

        doc.add_heading('4. 生產穩定性', level=1)
        doc.add_paragraph(clean_text_for_word(texts['stability_analysis']))
        add_chart_with_desc('cv', 'CV 變異係數', 'cv_desc')
        add_chart_with_desc('scatter', '效率能耗矩陣', 'scatter_desc')
        
        doc.add_heading('5. 策略行動建議', level=1)
        doc.add_paragraph(clean_text_for_word(texts['action_plan']))
        
        bio = BytesIO()
        doc.save(bio)
        return bio

# ==========================================
# 5. Main App
# ==========================================
def main():
    st.markdown("### 📥 數據輸入控制台")
    uploaded_file = st.file_uploader("匯入生產報表 (Excel/CSV)", type=["xlsx", "csv"], label_visibility="collapsed")
    
    if 'input_data' not in st.session_state:
        st.session_state.input_data = pd.DataFrame([
            {"日期": "2025-11-17", "廠別": "A廠", "機台編號": "ACO2", "OEE(%)": 50.1, "產量(雙)": 2009.5, "用電量(kWh)": 6.2},
            {"日期": "2025-11-17", "廠別": "A廠", "機台編號": "ACO4", "OEE(%)": 55.4, "產量(雙)": 4416.5, "用電量(kWh)": 9.1},
            {"日期": "2025-11-18", "廠別": "A廠", "機台編號": "ACO2", "OEE(%)": 48.5, "產量(雙)": 1950.0, "用電量(kWh)": 6.0},
        ])
        st.session_state.input_data['日期'] = pd.to_datetime(st.session_state.input_data['日期']).dt.date
    
    if uploaded_file:
        try:
            if uploaded_file.name.endswith('.csv'): df_new = pd.read_csv(uploaded_file)
            else: df_new = pd.read_excel(uploaded_file)
            rename_map = {"用電量(kWh)": "耗電量", "產量(雙)": "產量", "OEE(%)": "OEE_RAW", "設備": "機台編號", "機台": "機台編號"}
            for user_col, sys_col in rename_map.items():
                if user_col in df_new.columns: df_new = df_new.rename(columns={user_col: sys_col})
            st.session_state.input_data = df_new
        except: st.error("檔案讀取失敗")

    edited_df = st.data_editor(st.session_state.input_data, num_rows="dynamic", use_container_width=True)
    
    if st.button("🗑️ 清空所有數據"):
        st.session_state.input_data = pd.DataFrame(columns=["日期", "廠別", "機台編號", "OEE(%)", "產量(雙)", "用電量(kWh)"])
        st.rerun()

    st.markdown("---")
    st.markdown("#### ⚙️ 分析參數設定")
    c1, c2, c3 = st.columns(3)
    params = {
        'elec_price': c1.number_input("電價 (元/度)", value=3.5, step=0.1),
        'target_oee': c2.number_input("目標 OEE (%)", value=85.0, step=0.5),
        'product_margin': c3.number_input("獲利估算 (元/雙)", value=10.0, step=1.0)
    }
    
    st.write("")
    col_run, col_export = st.columns([1, 1])
    
    data_ready = False
    df_res, summary_res, scope_res, texts_res, figs_res = None, None, None, None, {}

    if not edited_df.empty:
        try:
            res = DataEngine.clean_and_process(edited_df, params)
            if res[0] is not None:
                df_res, summary_res, scope_res = res
                data_ready = True
                texts_res = InsightEngine.generate_narrative(df_res, summary_res, 
                                                           "廠別" if scope_res=="跨廠區分析" else "機台編號", 
                                                           params)
                figs_res = {
                    'rank': VizEngine.create_rank_chart(summary_res, "廠別" if scope_res=="跨廠區分析" else "機台編號"),
                    'cv': VizEngine.create_cv_chart(df_res, "廠別" if scope_res=="跨廠區分析" else "機台編號"),
                    'scatter': VizEngine.create_scatter_chart(df_res, "廠別" if scope_res=="跨廠區分析" else "機台編號"),
                    'dual': VizEngine.create_dual_axis_chart(df_res, "廠別" if scope_res=="跨廠區分析" else "機台編號"),
                    'pie': VizEngine.create_pie_chart(summary_res, "廠別" if scope_res=="跨廠區分析" else "機台編號"),
                    'unit': VizEngine.create_unit_energy_chart(summary_res, "廠別" if scope_res=="跨廠區分析" else "機台編號")
                }
            elif isinstance(res[2], str): st.warning(res[2])
        except Exception as e: st.error(f"數據處理錯誤: {e}")

    with col_run:
        start_btn = st.button("🚀 啟動全方位分析", type="primary")
        
    with col_export:
        if data_ready:
            try:
                docx = ReportEngine.generate_docx(df_res, summary_res, texts_res, figs_res, scope_res)
                st.download_button("📥 下載 Word 報告", docx.getvalue(), 
                                 f"生產效能報告_{pd.Timestamp.now().strftime('%Y%m%d')}.docx",
                                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e: st.error(f"匯出失敗: {e}")
        else: st.button("📥 下載 Word 報告", disabled=True)

    if start_btn and data_ready:
        with st.spinner('正在進行深度診斷...'):
            time.sleep(0.5)
            st.markdown("---")
            st.title("生產效能診斷分析報告")
            
            st.header("1. 總體績效概覽")
            st.markdown(f'<div class="insight-box">{md_to_html(texts_res["kpi_summary"])}</div>', unsafe_allow_html=True)
            st.subheader("績效總表")
            st.dataframe(summary_res.style.format({"OEE": "{:.1%}", "平均單位能耗": "{:.5f}", "總損失": "${:,.0f}"}).background_gradient(subset=["OEE"], cmap="Blues"), use_container_width=True)
            st.plotly_chart(figs_res['rank'], use_container_width=True)
            st.markdown(f'<div class="chart-desc">{md_to_html(texts_res["rank_desc"])}</div>', unsafe_allow_html=True)
            
            st.header("2. 深度診斷分析")
            st.markdown(f'<div class="analysis-text">{md_to_html(texts_res["benchmark_analysis"])}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="insight-box">{md_to_html(texts_res["opportunity_analysis"])}</div>', unsafe_allow_html=True)
            st.subheader("產量與能耗趨勢")
            st.plotly_chart(figs_res['dual'], use_container_width=True)
            st.markdown(f'<div class="chart-desc">{md_to_html(texts_res["dual_desc"])}</div>', unsafe_allow_html=True)
            
            st.header("3. 電力耗能深度分析")
            st.plotly_chart(figs_res['pie'], use_container_width=True)
            st.markdown(f'<div class="chart-desc">{md_to_html(texts_res["pie_desc"])}</div>', unsafe_allow_html=True)
            st.plotly_chart(figs_res['unit'], use_container_width=True)
            st.markdown(f'<div class="chart-desc">{md_to_html(texts_res["unit_desc"])}</div>', unsafe_allow_html=True)
            
            st.header("4. 生產趨勢與穩定性")
            st.plotly_chart(figs_res['cv'], use_container_width=True)
            st.markdown(f'<div class="chart-desc">{md_to_html(texts_res["cv_desc"])}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="analysis-text">{md_to_html(texts_res["stability_analysis"])}</div>', unsafe_allow_html=True)
            st.plotly_chart(figs_res['scatter'], use_container_width=True)
            st.markdown(f'<div class="chart-desc">{md_to_html(texts_res["scatter_desc"])}</div>', unsafe_allow_html=True)
            
            st.header("5. 綜合診斷與建議")
            st.markdown(texts_res['action_plan'])

if __name__ == "__main__":
    main()
